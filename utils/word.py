import os

from django.db.models import Count, Q
from docx import Document
from docx.shared import Pt

from Pdn.models import Person, Podrazdelenie


def word_skzi(request, skzis):
    doc = Document('static/files/obrazets_prilozhenie_14.1.docx')
    table = doc.tables[0]
    for i, skzi in enumerate(skzis):
        table.add_row()
        row = i + 3
        table.cell(row, 0).text = str(i + 1)
        table.cell(row, 1).text = str(skzi.name)
        table.cell(row, 2).text = skzi.serial_n
        table.cell(row, 3).text = skzi.ekz_n
        table.cell(row, 4).text = str(skzi.from_organ)
        l = []
        if skzi.date_poluch:
            l.append(skzi.date_poluch.strftime('%d.%m.%Y'))
        if skzi.nomer_poluch:
            l.append(skzi.nomer_poluch)
        table.cell(row, 5).text = ', '.join(l)

        table.cell(row, 6).text = skzi.to_person.get_fio()
        l = []
        if skzi.date_sopr:
            l.append(skzi.date_sopr.strftime('%d.%m.%Y'))
        if skzi.nomer_sopr:
            l.append(skzi.nomer_sopr)
        table.cell(row, 7).text = ', '.join(l)
        l = []
        if skzi.date_podtv:
            l.append(skzi.date_podtv.strftime('%d.%m.%Y'))
        if skzi.nomer_podtv:
            l.append(skzi.nomer_podtv)
        table.cell(row, 7).text = ', '.join(l)
        l = []
        if skzi.vozvr_date_sopr:
            l.append(skzi.vozvr_date_sopr.strftime('%d.%m.%Y'))
        if skzi.vozvr_nomer_sopr:
            l.append(skzi.vozvr_nomer_sopr)
        table.cell(row, 8).text = ', '.join(l)
        l = []
        if skzi.vozvr_date_podtv:
            l.append(skzi.vozvr_date_podtv.strftime('%d.%m.%Y'))
        if skzi.vozvr_nomer_podtv:
            l.append(skzi.vozvr_nomer_podtv)
        table.cell(row, 9).text = ', '.join(l)

        if skzi.date_vvod:
            table.cell(row, 10).text = skzi.date_vvod.strftime('%d.%m.%Y')
        if skzi.date_vivod:
            table.cell(row, 11).text = skzi.date_vivod.strftime('%d.%m.%Y')
        if skzi.date_unichtozh:
            table.cell(row, 12).text = skzi.date_unichtozh.strftime('%d.%m.%Y')
        if skzi.nomer_acta:
            table.cell(row, 13).text = skzi.nomer_acta
        if skzi.primechanie:
            table.cell(row, 14).text = skzi.primechanie

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(9)

    path = f'tmp/skzi/word/{request.user.id}'

    if not os.path.exists(path):
        os.makedirs(path)

    filename = f'Журнал СКЗИ.docx'
    path += f'/{filename}'

    doc.save(path)
    return path


def word_skzi_mini(request, skzis):
    doc = Document('static/files/Skzi-mini.docx')
    table = doc.tables[0]

    for i, skzi in enumerate(skzis):
        table.add_row()
        row = i + 1
        table.cell(row, 0).text = str(i + 1)
        table.cell(row, 1).text = str(skzi.name)
        table.cell(row, 2).text = skzi.serial_n
        table.cell(row, 3).text = skzi.to_person.get_fio()
        if skzi.date_vvod:
            table.cell(row, 4).text = skzi.date_vvod.strftime('%d.%m.%Y')

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(12)

    path = f'tmp/skzi/word/{request.user.id}'

    if not os.path.exists(path):
        os.makedirs(path)

    filename = f'Отчет СКЗИ.docx'
    path += f'/{filename}'

    doc.save(path)
    return path


def persons_otchet(request):
    podrazds = Podrazdelenie.objects.values('name'). \
        annotate(person_count=Count('persons'),
                 person_sogl=Count('persons',
                                   filter=Q(persons__soglasie=True) & Q(persons__sogl_raspr=True))).order_by('name')
    doc = Document('static/files/Persons_sogl.docx')
    table = doc.tables[0]

    for i, podrazd in enumerate(podrazds):
        table.add_row()
        row = i + 1
        table.cell(row, 0).text = str(i + 1)
        table.cell(row, 1).text = str(podrazd['name'])
        table.cell(row, 2).text = str(podrazd['person_count'])
        table.cell(row, 3).text = str(podrazd['person_sogl'])
        table.cell(row, 4).text = str(podrazd['person_count'] - podrazd['person_sogl'])

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(12)

    path = f'tmp/skzi/word/{request.user.id}'

    if not os.path.exists(path):
        os.makedirs(path)

    filename = f'Отчет ПДн.docx'
    path += f'/{filename}'

    doc.save(path)
    return path